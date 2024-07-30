from docx import Document
from docx.shared import Inches

def create_word_with_images(image1_path, image2_path, output_path):
    # Crear un nuevo documento
    doc = Document()
    
    # Configurar la orientación de la página a horizontal
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    # Calcular el tamaño de las imágenes para que quepan en la misma hoja
    max_width = new_width - Inches(1)  # Ajuste para márgenes
    image_height = Inches(2)  # Ajustar según sea necesario

    # Agregar la primera imagen centrada
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(image1_path, width=max_width, height=image_height)
    p.alignment = 1  # Alineación central

    # Agregar la segunda imagen centrada
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(image2_path, width=max_width, height=image_height)
    p.alignment = 1  # Alineación central

    # Guardar el documento
    doc.save(output_path)

# Ruta de las imágenes y del archivo de salida
image1_path = 'a1.jpg'
image2_path = 'tabla.jpg'
output_path = 'documento_con_imagenes.docx'

# Crear el documento
create_word_with_images(image1_path, image2_path, output_path)
